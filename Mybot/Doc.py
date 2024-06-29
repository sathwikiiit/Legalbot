from docx import Document
from docx.shared import Pt,Inches
from docx.enum.text import WD_LINE_SPACING,WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup, Tag
def create_document(html,path):
    '''
    margins here are 0.5 on all sides
    Times New Roman, font 16, 
    h1: bold center
    h2: not bold, works with style attribute(left,right,center) style:"text.align:left;"
    h3: indented not bold
    h4: underlined bold works with style attribute(left,center) style:"text.align:left;"
    ol: Numbered List unindented creates new level for each list
    p: simple para
    '''
    listlevel=0
    parser=BeautifulSoup(html,'html.parser')
    for x in parser.children: 
        if x.string=='\n':     # type: ignore # Remove empty name 
            x.extract()
    document = Document()
    section = document.sections[0]  # Assuming you want to modify the first section
    section.top_margin = Inches(0.5)  # 1 inch in EMUs
    section.right_margin = Inches(0.5)  # 1.5 inches in EMUs
    section.bottom_margin = Inches(0.5)  # 0.75 inches in EMUs
    section.left_margin = Inches(0.5)
    # Access the "Normal" style (optional, modify styles as needed)
    normal_style = document.styles['Normal']
    normal_style.font.name = "Times New Roman" # type: ignore
    normal_style.font.size = Pt(16) # type: ignore
    for i in parser.find_all():
        if str(i.name) == 'h1': # type: ignore
            x=document.add_paragraph(str(i.string)) # type: ignore
            x.runs[0].bold=True
            x.alignment=WD_ALIGN_PARAGRAPH.CENTER
        elif i.name == 'h2' and i.attrs['style']=='text-align: left;': # type: ignore
            x=document.add_paragraph(str(i.string)) # type: ignore
            x.runs[0].bold=False
            x.alignment=WD_ALIGN_PARAGRAPH.LEFT
        elif i.name == 'h2' and i.attrs['style']=='text-align: right;': # type: ignore
            x=document.add_paragraph(str(i.string)) # type: ignore
            x.runs[0].bold=False
            x.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        elif i.name == 'h2' and i.attrs['style']=='text-align: center;': # type: ignore
            x=document.add_paragraph(str(i.string)) # type: ignore
            x.runs[0].bold=False
            x.alignment=WD_ALIGN_PARAGRAPH.CENTER
        elif i.name == 'h3': # type: ignore
            x=document.add_paragraph(str(i.string)) # type: ignore
            x.runs[0].bold=False
            x.paragraph_format.left_indent = Inches(1)
            x.paragraph_format.right_indent= Inches(9)
        elif i.name == 'p':
            para=str(i.string).split('*') # type: ignore
            if len(para)>1:
                for i in range(len(para)):
                    x=document.add_paragraph(para[i]) # type: ignore
                    x.runs[i].bold=False if i%2==0 else True
                    x.runs[i].underline=False if i%2==0 else True
                    x.paragraph_format.left_indent = Inches(0)
                    x.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                x=document.add_paragraph(para[0]) # type: ignore
                x.runs[0].bold=False
                x.paragraph_format.left_indent = Inches(0)
                x.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        elif i.name == 'ol' and len(i.find_all())>1:
            for j in i.find_all():
                x=document.add_paragraph(str(j.string), style='List Number')
                x.paragraph_format.left_indent=Inches(0)
            listlevel+=1
        elif i.name == 'ol' and len(i.find_all())==1:
            for j in i.find_all():
                x=document.add_paragraph(str(j.string), style='Normal')
                x.paragraph_format.left_indent=Inches(0)
        elif i.name == 'h4' and i.attrs['style']=='text-align: center;': # type: ignore
            x=document.add_paragraph(str(i.string)) # type: ignore
            x.runs[0].bold=False
            x.runs[0].underline=True
            x.alignment=WD_ALIGN_PARAGRAPH.CENTER
        elif i.name == 'h4' and i.attrs['style']=='text-align: left;': # type: ignore
            x=document.add_paragraph(str(i.string)) # type: ignore
            x.runs[0].bold=False
            x.runs[0].underline=True
            x.alignment=WD_ALIGN_PARAGRAPH.CENTER
        elif i.name =='br':
            document.add_page_break()
        else:
            document.add_paragraph('Not recognised command {}'.format(str(i)))
    document.save(path)
def create_suit(html,path):
    '''
    margins here are 0.5 on all sides
    Times New Roman, font 16, 
    h1: bold center
    h2: not bold, works with style attribute(left,right,center) style:"text.align:left;"
    h3: indented not bold
    h4: underlined bold works with style attribute(left,center) style:"text.align:left;"
    ol: Numbered List unindented creates new level for each list
    p: simple para
    '''
    Alphabets= [chr(65+i) for i in range(26)]
    level=['',' 2',' 3']
    listlevel=0
    parser=BeautifulSoup(html,'html.parser')
    for x in parser.children: 
        if x.string=='\n':     # type: ignore # Remove empty name 
            x.extract()
    document = Document()
    section = document.sections[0]  # Assuming you want to modify the first section
    section.top_margin = Inches(1)  # 1 inch in EMUs
    section.right_margin = Inches(1)  # 1.5 inches in EMUs
    section.bottom_margin = Inches(1)  # 0.75 inches in EMUs
    section.left_margin = Inches(1)
    # Access the "Normal" style (optional, modify styles as needed)
    normal_style = document.styles['Normal']
    normal_style.font.name = "Times New Roman" # type: ignore
    normal_style.font.size = Pt(13) # type: ignore
    for i in parser.find_all():
        if str(i.name) == 'h1': # type: ignore
            x=document.add_paragraph(str(i.string)) # type: ignore
            x.runs[0].bold=True
            x.alignment=WD_ALIGN_PARAGRAPH.CENTER
        elif i.name == 'h2' and i.attrs['style']=='text-align: left;': # type: ignore
            x=document.add_paragraph(str(i.string)) # type: ignore
            x.runs[0].bold=False
            x.alignment=WD_ALIGN_PARAGRAPH.LEFT
        elif i.name == 'h2' and i.attrs['style']=='text-align: right;': # type: ignore
            x=document.add_paragraph(str(i.string)) # type: ignore
            x.runs[0].bold=False
            x.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        elif i.name == 'h2' and i.attrs['style']=='text-align: center;': # type: ignore
            x=document.add_paragraph(str(i.string)) # type: ignore
            x.runs[0].bold=False
            x.alignment=WD_ALIGN_PARAGRAPH.CENTER
        elif i.name == 'h3': # type: ignore
            x=document.add_paragraph(str(i.string)) # type: ignore
            x.runs[0].bold=False
            x.paragraph_format.left_indent = Inches(1)
            x.paragraph_format.right_indent= Inches(9)
        elif i.name == 'p':
            para=str(i.string).split('*') # type: ignore
            if len(para)>1:
                x=document.add_paragraph()
                for i in range(len(para)):
                    x.add_run(para[i]) # type: ignore
                    x.runs[i].bold=False if i%2==0 else True
                    x.runs[i].underline=False if i%2==0 else True
                    x.paragraph_format.left_indent = Inches(0)
                    x.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                x=document.add_paragraph(para[0]) # type: ignore
                x.runs[0].bold=False
                x.paragraph_format.left_indent = Inches(0)
                x.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        elif i.name == 'ol' and i.attrs['style']=='A':
            k=0
            for j in i.find_all():
                x=document.add_paragraph()
                x.add_run('{}.'.format(Alphabets[k])).add_tab()
                x.add_run(j.string+'\n')
                k+=1
        elif i.name == 'ol' and i.attrs['style']=='1':
            k=1
            for j in i.find_all():
                x=document.add_paragraph()
                x.add_run('{}.'.format(str(k))).add_tab()
                x.add_run(j.string+'\n')
                k+=1
        elif i.name == 'h4' and i.attrs['style']=='text-align: center;': # type: ignore
            x=document.add_paragraph(str(i.string)) # type: ignore
            x.runs[0].bold=False
            x.runs[0].underline=True
            x.alignment=WD_ALIGN_PARAGRAPH.CENTER
        elif i.name == 'h4' and i.attrs['style']=='text-align: left;': # type: ignore
            x=document.add_paragraph(str(i.string)) # type: ignore
            x.runs[0].bold=False
            x.runs[0].underline=True
            x.alignment=WD_ALIGN_PARAGRAPH.CENTER
        elif i.name =='br':
            document.add_page_break()
    document.save(path)

if __name__=='__main__':
    html ='''<h1> NOTICE TO SHOW CAUSE</h1>
    <h1>IN THE COURT OF CourtName</h1>
    <h1>AT City</h1>
    <h1>I.A. No: IANumber OF 2024</h1>
    <h1>IN</h1>
    <h1>O.S No: OSNumber OF 2024</h1>
    <h2 style="text-align: left;">BETWEEN:</h2>
    <h2 style="text-align: left;">Plaintiff1</h2>
    <h2 style="text-align: right;">tag1</h2>
    <h2 style="text-align: center;">AND</h2>
    <h2 style="text-align: left;">Defendant1</h2>
    <h2 style="text-align: right;">tag2</h2>
    <h3>To</h3>
    <h3>i</h3>
    <p>Whereas the above Named Petitioner_________ has made an Application to this court. U/O 31 Rule 1 & 2 of CPC</p>
    <p>You are hereby wanted to appear in this court in person, or by a pleader duly instructed on the ___ day of _________ 2024 at 10:30 a.m. in the forenoon to show Cause of against the Application, failing which the said Application will be heard and determined exparty. Given under my hand and seal of the Court this______day of ____________2024</p>
    <br>
'''

    create_document(html,'new.docx')
